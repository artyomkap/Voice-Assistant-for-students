import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import config

def create_word_file():
    doc = Document()

    first_paragraph = doc.add_paragraph('NOTTINGHAM TRENT UNIVERSITY\n SCHOOL OF SCIENCE AND TECHNOLOGY\n\n\n\n\n\n\n\n')
    first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = first_paragraph.runs[0]
    font = run.font
    font.size = Pt(18)
    font.name = 'Verdana'

    second_paragraph = doc.add_paragraph(f'Voice Assistant for Students\n by\n {config.STUDENT_NAME}\n in\n 2024\n\n\n\n\n\n\n')
    second_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    second_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run2 = second_paragraph.runs[0]
    font2 = run2.font
    font2.size = Pt(14)
    font2.name = 'Verdana'
    run2.bold = True

    third_paragraph = doc.add_paragraph(f'Project report in part fulfilment\n of the requirements for the degree of\n Bachelor of Science with Honours\n In\n {config.STUDENT_COURSE}')
    third_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    third_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run3 = third_paragraph.runs[0]
    font3 = run3.font
    font3.size = Pt(10)
    font3.name = 'Verdana'
    run3.bold = True


    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    file_path = os.path.join(desktop_path, 'New Coursework.docx')
    doc.save(file_path)



